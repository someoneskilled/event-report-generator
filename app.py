from flask import Flask, render_template, request, session, send_file
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import os
import base64
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from openpyxl import load_workbook

# === Setup ===
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    raise ValueError("GROQ_API_KEY not found in environment variables")
client = Groq(api_key=api_key)

app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'static/uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# === Utility Functions ===
def generateAI(input_text, instruction):
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": instruction},
                {"role": "user", "content": input_text}
            ],
            temperature=0.5,
            top_p=1
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {str(e)}"

def generate_image_caption(image_path):
    try:
        with open(image_path, "rb") as img_file:
            b64_image = base64.b64encode(img_file.read()).decode("utf-8")
        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Give a short one-liner tagline for the image of event, no intro, just few words caption telling what's in the image."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64_image}"}}
                    ]
                }
            ],
            temperature=1,
            max_tokens=1024,
            top_p=1,
            stream=False
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Caption Error: {str(e)}"

# === Routes ===
@app.route('/')
def form():
    return render_template('form.html')

@app.route('/submit', methods=['POST'])
def submit():
    form_data = {}
    text_fields = [
        "event_title", "web_url", "event_summary", "event_objectives", "event_feedback_bulletin",
        "event_venue", "event_date", "event_department",
        "faculty_coordinator", "chief_guest", "guest_testimonial"
    ]

    for field in text_fields:
        form_data[field] = request.form.get(field)

    if form_data["event_summary"]:
        form_data["event_summary"] = generateAI(
            form_data["event_summary"],
            "You are a professional event summarizer. Convert rough notes into a clear 150-word summary. Keep it professional, simple, and ready for documentation."
        )
    if form_data["event_objectives"]:
        form_data["event_objectives"] = generateAI(
            form_data["event_objectives"],
            "You are an academic writer. Convert these rough event objectives into a clear bullet-point list."
        )
    if form_data["event_feedback_bulletin"]:
        form_data["event_feedback_bulletin"] = generateAI(
            form_data["event_feedback_bulletin"],
            "Turn these rough notes into polished feedback points, no intro, only bullet points."
        )

    extra_input = f"""
    Summary: {form_data["event_summary"]}
    Objectives: {form_data["event_objectives"]}
    Activities: {form_data["event_feedback_bulletin"]}
    """
    instruction_extra = """
    Based on summary, objectives, and activities:
    1. Generate simple bullet points for Event Outcome.
    2. Generate SEO-Friendly Short Description.
    Format:
    Event Outcome: <points>
    SEO Description: <line>
    """
    extra_response = generateAI(extra_input.strip(), instruction_extra)

    try:
        parts = extra_response.split("SEO Description:")
        form_data["event_outcome"] = parts[0].replace("Event Outcome:", "").strip()
        form_data["seo_description"] = parts[1].strip()
    except Exception:
        form_data["event_outcome"] = "Error generating outcome."
        form_data["seo_description"] = "Error generating SEO description."

    form_data['event_time'] = request.form.getlist('event_time')

    # Handle files
    file_fields = {
        'poster': None,
        'certificates': [],
        'event_photos': [],
        'event_excel': None
    }

    # Poster
    poster = request.files.get('poster')
    if poster and poster.filename != '':
        filename = secure_filename(poster.filename)
        path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        poster.save(path)
        file_fields['poster'] = filename

    # Certificates and Photos
    for field in ['certificates', 'event_photos']:
        uploaded_files = request.files.getlist(field)
        paths = []
        for file in uploaded_files:
            if file and file.filename != '':
                filename = secure_filename(file.filename)
                path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(path)
                paths.append(filename)
        file_fields[field] = paths

    # Event Excel
    event_excel = request.files.get('event_excel')
    if event_excel and event_excel.filename.endswith(('.xls', '.xlsx')):
        excel_filename = secure_filename(event_excel.filename)
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], excel_filename)
        event_excel.save(excel_path)
        file_fields['event_excel'] = excel_filename

    # Image captions
    image_captions = {}
    for photo_filename in file_fields['event_photos']:
        image_path = os.path.join(app.config['UPLOAD_FOLDER'], photo_filename)
        caption = generate_image_caption(image_path)
        image_captions[photo_filename] = caption

    # Save to session
    session['form_data'] = form_data
    session['image_captions'] = image_captions
    session['file_fields'] = file_fields

    return render_template('summary.html', data={"form_data": form_data, "file_fields": file_fields, "image_captions": image_captions})

@app.route('/download-docx')
def download_docx():
    form_data = session.get('form_data')
    image_captions = session.get('image_captions', {})
    file_fields = session.get('file_fields', {})

    if not form_data:
        return "No data to export."

    doc = Document()

    # === Styling Utilities ===
    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_content(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_big_image(filename):
        try:
            path = os.path.join(app.root_path, 'static', 'uploads', filename)
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(5))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                add_content(f"[Poster not found: {filename}]")
        except Exception:
            add_content(f"[Error loading poster: {filename}]")

    def add_small_images_side_by_side(filenames):
        table = doc.add_table(rows=1, cols=len(filenames))
        table.autofit = True
        row = table.rows[0]
        for idx, filename in enumerate(filenames):
            cell = row.cells[idx]
            path = os.path.join(app.root_path, 'static', 'uploads', filename)
            if os.path.exists(path):
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(path, width=Inches(2))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption = image_captions.get(filename, "")
                p = cell.add_paragraph(caption)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.runs[0].italic = True
            else:
                cell.text = "[Image missing]"

    # === Document Content ===
    doc.add_heading('Kristu Jayanti College (Autonomous)', 0)

    ordered_fields = [
        ("Department", form_data.get("event_department")),
        ("Title", form_data.get("event_title")),
        ("Date", form_data.get("event_date")),
        ("Venue", form_data.get("event_venue")),
        ("Objectives", form_data.get("event_objectives")),
        ("Summary", form_data.get("event_summary")),
        ("Outcomes", form_data.get("event_outcome")),
        ("Web URL", form_data.get("web_url")),
        ("Feedback", form_data.get("event_feedback_bulletin")),
    ]

    for label, value in ordered_fields:
        if value:
            add_heading(label)
            add_content(value)

    if file_fields.get('poster'):
        add_heading("Poster")
        add_big_image(file_fields['poster'])

    if file_fields.get('event_photos'):
        add_heading("Event Photos")
        add_small_images_side_by_side(file_fields['event_photos'])

    if file_fields.get('certificates'):
        add_heading("Certificates")
        for cert in file_fields['certificates']:
            path = os.path.join(app.root_path, 'static', 'uploads', cert)
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(3))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # === Insert Excel as Table (Optimized version) ===
    if file_fields.get('event_excel'):
        excel_path = os.path.join(app.root_path, 'static', 'uploads', file_fields['event_excel'])
        wb = load_workbook(excel_path)
        sheet = wb.active

        # Collect only non-empty rows
        non_empty_rows = []
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None and str(cell).strip() != '' for cell in row):
                non_empty_rows.append(row)

        if non_empty_rows:
            # Find the maximum number of non-empty columns
            max_cols = max(len([cell for cell in row if cell is not None and str(cell).strip() != '']) for row in non_empty_rows)

            add_heading("Participants / Winners / Presenters List")
            table = doc.add_table(rows=1, cols=max_cols)
            table.style = 'Light List Accent 1'

            # Header
            hdr_cells = table.rows[0].cells
            for idx, cell in enumerate(non_empty_rows[0]):
                if idx < max_cols:
                    hdr_cells[idx].text = str(cell) if cell is not None else ''

            # Rows
            for row_data in non_empty_rows[1:]:
                row_cells = table.add_row().cells
                for idx, value in enumerate(row_data):
                    if idx < max_cols:
                        row_cells[idx].text = str(value) if value is not None else ''

    # === Save and Download ===
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name="summary.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True)
